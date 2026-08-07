from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


def set_run_font(run, ascii_font: str, east_asia_font: str, size_pt: float | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def add_paragraph_border(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F3")


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_style_fonts(style, ascii_font: str, east_asia_font: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = ascii_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    style.font.size = Pt(size_pt)
    style.font.bold = bold


def strip_markdown_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)


def write_inline_markdown(paragraph, text: str, base_size: float = 10.5) -> None:
    text = strip_markdown_links(text)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", "Consolas", base_size)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Calibri", "Microsoft YaHei", base_size)


def is_special_line(line: str) -> bool:
    stripped = line.strip()
    return (
        not stripped
        or stripped == "---"
        or stripped.startswith("```")
        or re.match(r"#{1,6}\s+", stripped) is not None
        or re.match(r"\s*[-*]\s+", line) is not None
        or re.match(r"\s*\d+\.\s+", line) is not None
    )


def list_style_name(kind: str, indent_level: int) -> str:
    level = min(indent_level, 2) + 1
    if kind == "bullet":
        return "List Bullet" if level == 1 else f"List Bullet {level}"
    return "List Number" if level == 1 else f"List Number {level}"


def collect_continuation(lines: list[str], start: int) -> tuple[str, int]:
    parts: list[str] = []
    i = start
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            break
        if re.match(r"\s{2,}.+", line) and not re.match(r"\s*[-*]\s+", line) and not re.match(r"\s*\d+\.\s+", line):
            parts.append(line.strip())
            i += 1
            continue
        break
    return (" ".join(parts), i)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)

    normal = doc.styles["Normal"]
    set_style_fonts(normal, "Calibri", "Microsoft YaHei", 10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    heading1 = doc.styles["Heading 1"]
    set_style_fonts(heading1, "Calibri", "Microsoft YaHei", 15, bold=True)
    heading1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = doc.styles["Heading 2"]
    set_style_fonts(heading2, "Calibri", "Microsoft YaHei", 12.5, bold=True)
    heading2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = doc.styles["Heading 3"]
    set_style_fonts(heading3, "Calibri", "Microsoft YaHei", 11, bold=True)
    heading3.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(3)

    for name in ["List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"]:
        if name in doc.styles:
            style = doc.styles[name]
            set_style_fonts(style, "Calibri", "Microsoft YaHei", 10.5)
            style.paragraph_format.space_after = Pt(3)


def add_title_block(doc: Document, title: str, subtitle: str) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, "Calibri", "Microsoft YaHei", 21)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    add_paragraph_border(title_paragraph)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(10)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    set_run_font(subtitle_run, "Calibri", "Microsoft YaHei", 10.5)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_code_block(doc: Document, code_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Mm(4)
    paragraph.paragraph_format.right_indent = Mm(4)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    shade_paragraph(paragraph, "F3F6FA")
    lines = code_text.splitlines()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, "Consolas", "Consolas", 9.5)
        if index != len(lines) - 1:
            run.add_break()


def markdown_to_docx(source_path: Path, output_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    title = "42688car 代码参考文档"
    subtitle = "当前代码对应的 CCS 智能小车模板参考文档。"
    for raw_line in lines:
        stripped = raw_line.strip()
        heading_match = re.match(r"#\s+(.*)", stripped)
        if heading_match:
            title = heading_match.group(1).strip()
            break

    add_title_block(doc, title, subtitle)

    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        heading_match = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            if level == 1:
                i += 1
                continue
            paragraph = doc.add_paragraph(style=f"Heading {min(level - 1, 3)}")
            write_inline_markdown(paragraph, heading_text, 12.5 if level <= 3 else 11)
            i += 1
            continue

        if stripped == "---":
            separator = doc.add_paragraph()
            separator.paragraph_format.space_after = Pt(8)
            add_paragraph_border(separator)
            i += 1
            continue

        bullet_match = re.match(r"(\s*)[-*]\s+(.*)", line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text_body = bullet_match.group(2).strip()
            continuation, next_index = collect_continuation(lines, i + 1)
            if continuation:
                text_body = f"{text_body} {continuation}"
            paragraph = doc.add_paragraph(style=list_style_name("bullet", indent_level))
            write_inline_markdown(paragraph, text_body)
            i = next_index if continuation else i + 1
            continue

        number_match = re.match(r"(\s*)\d+\.\s+(.*)", line)
        if number_match:
            indent_level = len(number_match.group(1)) // 2
            text_body = number_match.group(2).strip()
            continuation, next_index = collect_continuation(lines, i + 1)
            if continuation:
                text_body = f"{text_body} {continuation}"
            paragraph = doc.add_paragraph(style=list_style_name("number", indent_level))
            write_inline_markdown(paragraph, text_body)
            i = next_index if continuation else i + 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines) and not is_special_line(lines[i]):
            paragraph_lines.append(lines[i].strip())
            i += 1
        paragraph = doc.add_paragraph()
        write_inline_markdown(paragraph, " ".join(paragraph_lines))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a DOCX reference guide from CODE_EXPLANATION.md.")
    parser.add_argument(
        "--source",
        default="CODE_EXPLANATION.md",
        help="UTF-8 markdown source file.",
    )
    parser.add_argument(
        "--output",
        default="42688car_代码说明_最终.docx",
        help="Output DOCX path.",
    )
    args = parser.parse_args()

    source_path = Path(args.source).resolve()
    output_path = Path(args.output).resolve()
    markdown_to_docx(source_path, output_path)
    print(output_path)


if __name__ == "__main__":
    main()
